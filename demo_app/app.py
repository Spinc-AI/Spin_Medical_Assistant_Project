"""Desktop demo client for STT, Core_LLM, and the Orchestrator.

One tab per service. Every model selector offers a **connection-type choice**:
"Remote Local Model" (talks to our deployed service, as before) or
"Custom Cloud API" (talks straight to an OpenAI-compatible cloud API from
THIS app — no dependency on our servers at all for that call).

On the STT and Core_LLM tabs the choice sits above Host/Port, since it decides
whether Host/Port even matters. The Orchestrator always needs its own server
(coordinating a multi-step instruction is literally its job), so instead the
choice sits separately before each of its two models (STT / LLM) — they can
even use different cloud providers/keys, since STT and LLM external calls are
fully independent.

The Orchestrator tab's input/output layout is built dynamically from
GET /instructions/{id} (input.accepts, output.type), so it adapts
automatically as new instructions are added.

Run:  python app.py
(Needs: pip install -r requirements.txt. On Linux, tkinter itself is a system
package: sudo apt install python3-tk)
"""
import json
import os
import queue
import re
import tempfile
import threading
import wave
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

try:
    import sounddevice as sd
    MIC_ERROR = None
except Exception as exc:  # missing PortAudio / no audio device — keep the app usable
    sd = None
    MIC_ERROR = exc

DEFAULT_HOST = "localhost"
TIMEOUT_SHORT = 15
TIMEOUT_LOAD = 900   # model loading / inference can be slow -- must exceed Orchestrator's own
                      # internal HTTP_TIMEOUT (see Orchestrator/.env), or this becomes the
                      # bottleneck instead: a big local model (e.g. aya-expanse:32b) can take
                      # several minutes to reconcile a BuAli report, and Orchestrator's timeout
                      # to Core_LLM needs raising to match (default 120s is too short for that).
AUDIO_FILETYPES = [("Audio files", "*.wav *.mp3 *.flac *.ogg *.m4a"), ("All files", "*.*")]
MIC_SAMPLE_RATE = 16000  # matches the STT service's target rate

# The two connection-type choices offered by every model selector.
REMOTE_LOCAL_LABEL = "Remote Local Model"
CUSTOM_CLOUD_LABEL = "Custom Cloud API"
DEFAULT_CLOUD_BASE_URL = "https://api.openai.com/v1"
DEFAULT_CLOUD_STT_MODEL = "whisper-1"
DEFAULT_CLOUD_LLM_MODEL = "gpt-4o-mini"
FALLBACK_LANGUAGES = {"fa": "Persian", "en": "English"}  # used before/without a server round-trip
# Core_LLM's registered local models (Core_LLM/deployment/model.py) -- all served
# directly via transformers, not Ollama. One model is loaded at a time; this
# list covers the "Remote Local Model" dropdown on the Core_LLM tab.
LOCAL_LLM_MODELS = ["aya-expanse-8b", "aya-expanse-32b", "gemma-4-31b",
                    "gemma-4-e4b", "gemma-4-12b", "qwen3-omni-30b"]


def cloud_transcribe(base_url, api_key, model, audio_bytes, filename, language=None):
    """Transcribe directly against an OpenAI-compatible /audio/transcriptions endpoint."""
    data = {"model": model or DEFAULT_CLOUD_STT_MODEL}
    if language:
        data["language"] = language
    r = requests.post(
        f"{base_url.rstrip('/')}/audio/transcriptions",
        headers={"Authorization": f"Bearer {api_key}"},
        files={"file": (filename, audio_bytes)},
        data=data,
        timeout=TIMEOUT_LOAD,
    )
    r.raise_for_status()
    return r.json()["text"]


def cloud_chat(base_url, api_key, model, messages):
    """Chat directly against an OpenAI-compatible /chat/completions endpoint."""
    r = requests.post(
        f"{base_url.rstrip('/')}/chat/completions",
        headers={"Authorization": f"Bearer {api_key}"},
        json={"model": model or DEFAULT_CLOUD_LLM_MODEL, "messages": messages},
        timeout=TIMEOUT_LOAD,
    )
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"]


def run_bg(fn, *args, **kwargs):
    """Fire fn(*args, **kwargs) on a background thread so the GUI never freezes."""
    threading.Thread(target=fn, args=args, kwargs=kwargs, daemon=True).start()


class ModeSelector(ttk.Frame):
    """Label + dropdown choosing "Remote Local Model" vs "Custom Cloud API"."""

    def __init__(self, parent, label_text):
        super().__init__(parent)
        ttk.Label(self, text=label_text).grid(row=0, column=0, sticky="w")
        self.mode = tk.StringVar(value=REMOTE_LOCAL_LABEL)
        self.box = ttk.Combobox(self, textvariable=self.mode, width=20, state="readonly",
                                values=[REMOTE_LOCAL_LABEL, CUSTOM_CLOUD_LABEL])
        self.box.grid(row=0, column=1, sticky="w", padx=(6, 0))

    def is_cloud(self):
        return self.mode.get() == CUSTOM_CLOUD_LABEL


class CloudFieldsFrame(ttk.Frame):
    """Cloud model / API key / base URL fields. Grid this frame as one unit
    and show/hide it as a whole via grid()/grid_remove()."""

    def __init__(self, parent, default_model):
        super().__init__(parent)
        ttk.Label(self, text="Cloud model:").grid(row=0, column=0, sticky="w")
        self.model = tk.StringVar(value=default_model)
        ttk.Entry(self, textvariable=self.model, width=18).grid(row=0, column=1, sticky="w", padx=(0, 12))

        ttk.Label(self, text="API key:").grid(row=0, column=2, sticky="w")
        self.api_key = tk.StringVar()
        ttk.Entry(self, textvariable=self.api_key, show="*", width=24).grid(row=0, column=3, sticky="w")

        ttk.Label(self, text="Base URL:").grid(row=1, column=0, sticky="w", pady=(4, 0))
        self.base_url = tk.StringVar(value=DEFAULT_CLOUD_BASE_URL)
        ttk.Entry(self, textvariable=self.base_url, width=40).grid(
            row=1, column=1, columnspan=3, sticky="w", pady=(4, 0))


class MicRecorder:
    """Records mono 16-bit PCM from the default microphone until stopped."""

    def __init__(self, samplerate=MIC_SAMPLE_RATE, channels=1):
        self.samplerate = samplerate
        self.channels = channels
        self._queue = queue.Queue()
        self._stream = None

    def _callback(self, indata, frames, time_info, status):
        self._queue.put(bytes(indata))

    def start(self):
        self._queue = queue.Queue()
        self._stream = sd.RawInputStream(
            samplerate=self.samplerate, channels=self.channels,
            dtype="int16", callback=self._callback,
        )
        self._stream.start()

    def stop_and_save(self):
        """Stop recording and write the captured audio to a temp .wav file. Returns its path."""
        self._stream.stop()
        self._stream.close()
        chunks = []
        while not self._queue.empty():
            chunks.append(self._queue.get())
        path = tempfile.NamedTemporaryFile(suffix=".wav", delete=False).name
        with wave.open(path, "wb") as wf:
            wf.setnchannels(self.channels)
            wf.setsampwidth(2)  # int16 -> 2 bytes
            wf.setframerate(self.samplerate)
            wf.writeframes(b"".join(chunks))
        return path


class RecordButton(ttk.Button):
    """Toggle button: records from the mic, writes a temp WAV, and stores its
    path in `file_path_var` — the same variable a Browse button would set."""

    def __init__(self, parent, file_path_var):
        super().__init__(parent, text="Record mic", command=self._toggle)
        self.file_path_var = file_path_var
        self._recorder = None

    def _toggle(self):
        if self._recorder is None:
            self._start()
        else:
            self._stop()

    def _start(self):
        if sd is None:
            messagebox.showerror("Microphone", f"Microphone support unavailable: {MIC_ERROR}")
            return
        try:
            self._recorder = MicRecorder()
            self._recorder.start()
        except Exception as exc:
            messagebox.showerror("Microphone", f"Could not start recording: {exc}")
            self._recorder = None
            return
        self.config(text="Stop recording")

    def _stop(self):
        try:
            path = self._recorder.stop_and_save()
            self.file_path_var.set(path)
        except Exception as exc:
            messagebox.showerror("Microphone", f"Could not save recording: {exc}")
        finally:
            self._recorder = None
            self.config(text="Record mic")


class ConnectionBar(ttk.Frame):
    """Host/Port fields + a Check button + a colored connected/disconnected indicator."""

    def __init__(self, parent, default_port, health_path="/"):
        super().__init__(parent)
        self.health_path = health_path

        ttk.Label(self, text="Host:").grid(row=0, column=0, padx=(0, 4))
        self.host = tk.StringVar(value=DEFAULT_HOST)
        ttk.Entry(self, textvariable=self.host, width=16).grid(row=0, column=1, padx=(0, 8))

        ttk.Label(self, text="Port:").grid(row=0, column=2, padx=(0, 4))
        self.port = tk.StringVar(value=str(default_port))
        ttk.Entry(self, textvariable=self.port, width=6).grid(row=0, column=3, padx=(0, 8))

        ttk.Button(self, text="Check connection", command=self.check).grid(row=0, column=4, padx=(0, 8))
        self.indicator = ttk.Label(self, text="● unknown", foreground="gray")
        self.indicator.grid(row=0, column=5)

    @property
    def base_url(self):
        return f"http://{self.host.get().strip()}:{self.port.get().strip()}"

    def check(self):
        self._set("● checking...", "gray")
        run_bg(self._check_bg)

    def _check_bg(self):
        try:
            ok = requests.get(self.base_url + self.health_path, timeout=TIMEOUT_SHORT).status_code == 200
        except requests.RequestException:
            ok = False
        self.after(0, self._set, ("● connected" if ok else "● unreachable"),
                   ("green" if ok else "red"))

    def _set(self, text, color):
        self.indicator.config(text=text, foreground=color)


def error_detail(exc):
    """Best-effort extraction of a JSON {"detail": ...} body from a requests error."""
    resp = getattr(exc, "response", None)
    if resp is not None:
        try:
            return resp.json().get("detail", str(exc))
        except ValueError:
            pass
    return str(exc)


# ---------------------------------------------------------------------------
# Transcript -> Word (.docx) export
# ---------------------------------------------------------------------------
# Persian/Arabic-block characters. A line containing any of these is treated
# as an RTL paragraph; Word's own bidi algorithm still renders embedded Latin
# words/numbers within it left-to-right, so mixed Persian/English lines are
# fine without per-run splitting.
_RTL_CHAR_RE = re.compile(r"[\u0590-\u08FF\uFB1D-\uFDFF\uFE70-\uFEFF]")

# Characters illegal in Windows filenames, plus "/" and ":" which show up in
# HF model ids (e.g. "MohammadReza-Halakoo/persian-whisper-large-v3...") and
# in our "openai:gpt-4o-mini" convention.
_FILENAME_UNSAFE_RE = re.compile(r'[\\/:*?"<>|]')


def _looks_rtl(text):
    return bool(_RTL_CHAR_RE.search(text))


def sanitize_filename_part(name):
    """Turn a model id/name into something safe to put in a filename."""
    if not name:
        return ""
    name = name[len("openai:"):] if name.startswith("openai:") else name
    name = _FILENAME_UNSAFE_RE.sub("-", name)
    return name.strip("-_")


def build_transcript_filename(audio_path, model_parts, ext=".docx"):
    """<audio-stem>_<model1>_<model2>_..._<modelN>.ext, skipping empty parts."""
    stem = os.path.splitext(os.path.basename(audio_path))[0] if audio_path else "transcript"
    parts = [stem] + [sanitize_filename_part(m) for m in model_parts if m]
    return "_".join(p for p in parts if p) + ext


def save_text_as_docx(text, path):
    """Write `text` to a .docx at `path`, one paragraph per line. Each line's
    direction (RTL/LTR) is auto-detected so mixed Persian/English transcripts
    render correctly in Word. python-docx has no high-level RTL API, so the
    <w:bidi>/<w:rtl> elements are added directly.
    """
    doc = Document()
    for line in (text or "").split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        if _looks_rtl(line):
            p._p.get_or_add_pPr().append(OxmlElement("w:bidi"))
            run._r.get_or_add_rPr().append(OxmlElement("w:rtl"))
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.save(path)


class TranscriptSaver(ttk.Frame):
    """'Save to: <dir>  [Browse...]  [Save Transcript (.docx)]' row.

    Defaults the save directory to the source audio's own directory; a
    Browse... button lets the user override it (the override sticks until
    the app restarts). `get_text_and_parts()` is supplied by the owning tab
    and returns (text_to_save, audio_path, [model names in filename order]).
    """

    def __init__(self, parent, get_text_and_parts):
        super().__init__(parent)
        self._get_text_and_parts = get_text_and_parts
        self._overridden = False
        self.save_dir = tk.StringVar()

        ttk.Label(self, text="Save to:").grid(row=0, column=0, sticky="w")
        ttk.Entry(self, textvariable=self.save_dir, width=40, state="readonly").grid(
            row=0, column=1, sticky="w")
        ttk.Button(self, text="Browse...", command=self._browse).grid(row=0, column=2, sticky="w", padx=4)
        ttk.Button(self, text="Save Transcript (.docx)", command=self._save).grid(
            row=0, column=3, sticky="w", padx=(8, 0))

    def note_audio_path(self, audio_path):
        """Called by the owning tab whenever a new audio file is selected/used."""
        if audio_path and not self._overridden:
            self.save_dir.set(os.path.dirname(os.path.abspath(audio_path)))

    def _browse(self):
        d = filedialog.askdirectory(title="Choose where to save the transcript")
        if d:
            self.save_dir.set(d)
            self._overridden = True

    def _save(self):
        text, audio_path, model_parts = self._get_text_and_parts()
        if not text or not text.strip():
            messagebox.showwarning("Save Transcript", "Nothing to save yet — run a transcription first.")
            return
        save_dir = self.save_dir.get().strip() or (
            os.path.dirname(os.path.abspath(audio_path)) if audio_path else os.getcwd()
        )
        filename = build_transcript_filename(audio_path, model_parts)
        path = os.path.join(save_dir, filename)
        try:
            os.makedirs(save_dir, exist_ok=True)
            save_text_as_docx(text, path)
            messagebox.showinfo("Save Transcript", f"Saved to:\n{path}")
        except OSError as exc:
            messagebox.showerror("Save Transcript", f"Could not save: {exc}")


class OutputBox(scrolledtext.ScrolledText):
    """A read-only text area used for API output."""

    def __init__(self, parent, height=12):
        super().__init__(parent, width=80, height=height, state="disabled", wrap="word")

    def write(self, text):
        self.config(state="normal")
        self.delete("1.0", tk.END)
        self.insert(tk.END, text)
        self.config(state="disabled")


class ScrollableFrame(ttk.Frame):
    """A frame with a vertical scrollbar. Put content in `.body`, not on
    `self` directly. Used to wrap each notebook tab so content that doesn't
    fit the window (e.g. the Save Transcript button, below several stacked
    STT slot widgets) is still reachable by scrolling instead of requiring
    the window to grow to fit everything at once."""

    def __init__(self, parent):
        super().__init__(parent)
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)

        canvas = tk.Canvas(self, highlightthickness=0)
        vsb = ttk.Scrollbar(self, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        canvas.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")

        self.body = ttk.Frame(canvas, padding=10)
        window = canvas.create_window((0, 0), window=self.body, anchor="nw")

        self.body.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind("<Configure>", lambda e: canvas.itemconfig(window, width=e.width))

        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        # Only scroll this canvas while the mouse is actually over it -- a
        # global bind_all would hijack scrolling everywhere (e.g. inside a
        # combobox's own dropdown list) regardless of which tab is active.
        canvas.bind("<Enter>", lambda e: canvas.bind_all("<MouseWheel>", _on_mousewheel))
        canvas.bind("<Leave>", lambda e: canvas.unbind_all("<MouseWheel>"))


# ---------------------------------------------------------------------------
# STT tab
# ---------------------------------------------------------------------------
class STTTab(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent, padding=10)
        self._local_models = []

        self.mode = ModeSelector(self, "Connection type:")
        self.mode.grid(row=0, column=0, columnspan=5, sticky="w", pady=(0, 8))
        self.mode.mode.trace_add("write", lambda *a: self._update_mode_visibility())

        # --- Remote Local Model group: Host/Port + local model controls ---
        self.local_frame = ttk.Frame(self)
        self.conn = ConnectionBar(self.local_frame, default_port=8000, health_path="/models")
        self.conn.grid(row=0, column=0, columnspan=5, sticky="w", pady=(0, 8))

        ttk.Label(self.local_frame, text="Model:").grid(row=1, column=0, sticky="w")
        self.model = tk.StringVar()
        self.model_box = ttk.Combobox(self.local_frame, textvariable=self.model, width=18, state="readonly")
        self.model_box.grid(row=1, column=1, sticky="w")
        ttk.Button(self.local_frame, text="Refresh", command=self.refresh_models).grid(
            row=1, column=2, sticky="w", padx=4)
        ttk.Button(self.local_frame, text="Load", command=self.load_model).grid(row=1, column=3, sticky="w")
        ttk.Button(self.local_frame, text="Unload", command=self.unload_model).grid(
            row=1, column=4, sticky="w", padx=4)
        self.model_status = ttk.Label(self.local_frame, text="loaded: -")
        self.model_status.grid(row=2, column=0, columnspan=3, sticky="w", pady=(2, 0))

        # --- Custom Cloud API group ---
        self.cloud = CloudFieldsFrame(self, DEFAULT_CLOUD_STT_MODEL)

        # --- Always visible: language + audio input + action + output ---
        ttk.Label(self, text="Language:").grid(row=2, column=0, sticky="w", pady=(8, 4))
        self.language = tk.StringVar()
        self.language_box = ttk.Combobox(self, textvariable=self.language, width=14, state="readonly",
                                         values=[f"{c} - {n}" for c, n in FALLBACK_LANGUAGES.items()])
        self.language_box.current(0)
        self.language_box.grid(row=2, column=1, sticky="w", pady=(8, 4))
        self._language_codes = list(FALLBACK_LANGUAGES.keys())

        ttk.Label(self, text="Audio file:").grid(row=3, column=0, sticky="w")
        self.file_path = tk.StringVar()
        ttk.Entry(self, textvariable=self.file_path, width=45, state="readonly").grid(
            row=3, column=1, columnspan=2, sticky="w")
        ttk.Button(self, text="Browse...", command=self.browse).grid(row=3, column=3, sticky="w")
        RecordButton(self, self.file_path).grid(row=3, column=4, sticky="w", padx=4)
        ttk.Button(self, text="Transcribe", command=self.transcribe).grid(row=4, column=0, sticky="w", pady=10)

        ttk.Label(self, text="Output:").grid(row=5, column=0, sticky="nw")
        self.output = OutputBox(self)
        self.output.grid(row=6, column=0, columnspan=5, pady=4)

        self._last_audio_path = ""
        self._last_stt_model = ""
        self.saver = TranscriptSaver(self, self._transcript_parts)
        self.saver.grid(row=7, column=0, columnspan=5, sticky="w", pady=(4, 0))

        self._update_mode_visibility()
        self.refresh_models()
        self.refresh_languages()

    def _transcript_parts(self):
        return self.output.get("1.0", tk.END).strip(), self._last_audio_path, [self._last_stt_model]

    def _is_cloud(self):
        return self.mode.is_cloud()

    def _update_mode_visibility(self):
        """Remote Local Model: Host/Port + local model controls. Custom Cloud API: the
        opposite — and that path never touches Host/Port/our server at all."""
        if self._is_cloud():
            self.local_frame.grid_remove()
            self.cloud.grid(row=1, column=0, columnspan=5, sticky="w", pady=(0, 8))
        else:
            self.cloud.grid_remove()
            self.local_frame.grid(row=1, column=0, columnspan=5, sticky="w", pady=(0, 8))

    def refresh_models(self):
        run_bg(self._refresh_models_bg)

    def _refresh_models_bg(self):
        try:
            r = requests.get(f"{self.conn.base_url}/models", timeout=TIMEOUT_SHORT)
            r.raise_for_status()
            self._local_models = r.json().get("available", [])
            self.after(0, self._apply_models, r.json())
        except requests.RequestException as exc:
            self.after(0, messagebox.showerror, "STT", f"Could not fetch models: {error_detail(exc)}")

    def _apply_models(self, data):
        self.model_box["values"] = self._local_models
        loaded = data.get("loaded")
        self.model_status.config(text=f"loaded: {loaded or '-'}")
        if loaded:
            self.model.set(loaded)
        elif self._local_models:
            self.model_box.current(0)

    def refresh_languages(self):
        run_bg(self._refresh_languages_bg)

    def _refresh_languages_bg(self):
        try:
            r = requests.get(f"{self.conn.base_url}/languages", timeout=TIMEOUT_SHORT)
            r.raise_for_status()
            self.after(0, self._apply_languages, r.json())
        except requests.RequestException:
            pass  # keep the FALLBACK_LANGUAGES already seeded — no server needed for this

    def _apply_languages(self, data):
        available = data.get("available", {})
        if not available:
            return
        self._language_codes = list(available.keys())
        self.language_box["values"] = [f"{c} - {n}" for c, n in available.items()]
        default = data.get("default")
        idx = self._language_codes.index(default) if default in self._language_codes else 0
        self.language_box.current(idx)

    def _selected_language(self):
        idx = self.language_box.current()
        return self._language_codes[idx] if 0 <= idx < len(self._language_codes) else None

    def load_model(self):
        key = self.model.get()
        if not key:
            messagebox.showwarning("STT", "Pick a model first.")
            return
        self.model_status.config(text=f"loading {key}...")
        run_bg(self._load_model_bg, key)

    def _load_model_bg(self, key):
        try:
            r = requests.post(f"{self.conn.base_url}/models/{key}/load", timeout=TIMEOUT_LOAD)
            r.raise_for_status()
            self.after(0, self.model_status.config, {"text": f"loaded: {r.json().get('model')}"})
        except requests.RequestException as exc:
            self.after(0, messagebox.showerror, "STT", f"Load failed: {error_detail(exc)}")
            self.after(0, self.model_status.config, {"text": "loaded: -"})

    def unload_model(self):
        run_bg(self._unload_model_bg)

    def _unload_model_bg(self):
        try:
            r = requests.post(f"{self.conn.base_url}/models/unload", timeout=TIMEOUT_SHORT)
            r.raise_for_status()
            self.after(0, self.model_status.config, {"text": "loaded: -"})
        except requests.RequestException as exc:
            self.after(0, messagebox.showerror, "STT", f"Unload failed: {error_detail(exc)}")

    def browse(self):
        path = filedialog.askopenfilename(title="Choose an audio file", filetypes=AUDIO_FILETYPES)
        if path:
            self.file_path.set(path)

    def transcribe(self):
        path = self.file_path.get()
        if not path:
            messagebox.showwarning("STT", "Choose an audio file first.")
            return
        language = self._selected_language()
        self.output.write("transcribing...")
        self._last_audio_path = path
        self.saver.note_audio_path(path)
        if self._is_cloud():
            self._last_stt_model = "openai:" + self.cloud.model.get().strip()
            run_bg(self._transcribe_cloud_bg, path, language)
        else:
            self._last_stt_model = self.model.get().strip()
            run_bg(self._transcribe_bg, path, language)

    def _transcribe_bg(self, path, language):
        try:
            with open(path, "rb") as f:
                data = {"language": language} if language else None
                r = requests.post(f"{self.conn.base_url}/transcribe",
                                  files={"file": (os.path.basename(path), f)}, data=data, timeout=TIMEOUT_LOAD)
            r.raise_for_status()
            self.after(0, self.output.write, r.json().get("text", ""))
        except requests.RequestException as exc:
            self.after(0, self.output.write, f"Error: {error_detail(exc)}")

    def _transcribe_cloud_bg(self, path, language):
        """Calls the cloud API directly — no Host/Port, no our-server dependency."""
        api_key = self.cloud.api_key.get().strip()
        if not api_key:
            self.after(0, self.output.write, "Error: enter an API key first.")
            return
        try:
            with open(path, "rb") as f:
                text = cloud_transcribe(self.cloud.base_url.get().strip(), api_key,
                                        self.cloud.model.get().strip(), f.read(),
                                        os.path.basename(path), language)
            self.after(0, self.output.write, text)
        except requests.RequestException as exc:
            self.after(0, self.output.write, f"Error: {error_detail(exc)}")


# ---------------------------------------------------------------------------
# Core_LLM tab
# ---------------------------------------------------------------------------
class LLMTab(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent, padding=10)

        self.mode = ModeSelector(self, "Connection type:")
        self.mode.grid(row=0, column=0, columnspan=3, sticky="w", pady=(0, 8))
        self.mode.mode.trace_add("write", lambda *a: self._update_mode_visibility())

        # --- Remote Local Model group ---
        self.local_frame = ttk.Frame(self)
        self.conn = ConnectionBar(self.local_frame, default_port=8001, health_path="/")
        self.conn.grid(row=0, column=0, columnspan=3, sticky="w", pady=(0, 8))
        ttk.Label(self.local_frame, text="Model:").grid(row=1, column=0, sticky="w")
        self.model = tk.StringVar(value=LOCAL_LLM_MODELS[0])
        ttk.Combobox(self.local_frame, textvariable=self.model, width=20,
                    values=LOCAL_LLM_MODELS).grid(row=1, column=1, sticky="w")
        self.unload_btn = ttk.Button(self.local_frame, text="Unload model", command=self.unload_model)
        self.unload_btn.grid(row=1, column=2, sticky="w", padx=4)

        # --- Custom Cloud API group ---
        self.cloud = CloudFieldsFrame(self, DEFAULT_CLOUD_LLM_MODEL)

        ttk.Label(self, text="Message:").grid(row=2, column=0, sticky="nw", pady=(10, 0))
        self.input = scrolledtext.ScrolledText(self, width=70, height=6, wrap="word")
        self.input.grid(row=2, column=1, columnspan=2, pady=(10, 0), sticky="w")

        ttk.Button(self, text="Send", command=self.send).grid(row=3, column=0, sticky="w", pady=10)

        ttk.Label(self, text="Output:").grid(row=4, column=0, sticky="nw")
        self.output = OutputBox(self)
        self.output.grid(row=5, column=0, columnspan=3, pady=4)

        self._update_mode_visibility()

    def _is_cloud(self):
        return self.mode.is_cloud()

    def _update_mode_visibility(self):
        if self._is_cloud():
            self.local_frame.grid_remove()
            self.cloud.grid(row=1, column=0, columnspan=3, sticky="w", pady=(0, 8))
        else:
            self.cloud.grid_remove()
            self.local_frame.grid(row=1, column=0, columnspan=3, sticky="w", pady=(0, 8))

    def send(self):
        text = self.input.get("1.0", tk.END).strip()
        if not text:
            messagebox.showwarning("Core_LLM", "Type a message first.")
            return
        self.output.write("thinking...")
        if self._is_cloud():
            run_bg(self._send_cloud_bg, text)
        else:
            run_bg(self._send_bg, text)

    def _send_bg(self, text):
        payload = {"messages": [{"role": "user", "content": text}]}
        model = self.model.get().strip()
        if model:
            payload["model"] = model
        try:
            r = requests.post(f"{self.conn.base_url}/chat", json=payload, timeout=TIMEOUT_LOAD)
            r.raise_for_status()
            self.after(0, self.output.write, r.json().get("reply", ""))
        except requests.RequestException as exc:
            self.after(0, self.output.write, f"Error: {error_detail(exc)}")

    def _send_cloud_bg(self, text):
        """Calls the cloud API directly — no Host/Port, no our-server dependency."""
        api_key = self.cloud.api_key.get().strip()
        if not api_key:
            self.after(0, self.output.write, "Error: enter an API key first.")
            return
        try:
            reply = cloud_chat(self.cloud.base_url.get().strip(), api_key, self.cloud.model.get().strip(),
                               [{"role": "user", "content": text}])
            self.after(0, self.output.write, reply)
        except requests.RequestException as exc:
            self.after(0, self.output.write, f"Error: {error_detail(exc)}")

    def unload_model(self):
        run_bg(self._unload_model_bg)

    def _unload_model_bg(self):
        model = self.model.get().strip() or None
        try:
            r = requests.post(f"{self.conn.base_url}/unload",
                              params=({"model": model} if model else {}), timeout=TIMEOUT_SHORT)
            r.raise_for_status()
            self.after(0, messagebox.showinfo, "Core_LLM", "Model unloaded.")
        except requests.RequestException as exc:
            self.after(0, messagebox.showerror, "Core_LLM", f"Unload failed: {error_detail(exc)}")


# ---------------------------------------------------------------------------
# Orchestrator tab
# ---------------------------------------------------------------------------
class OrchestratorTab(ttk.Frame):
    """A chat client for the orchestrator's pipeline API: pick a pipeline,
    talk to it, watch it hand back a reply and — once it has one — a
    structured result. No model fields anywhere here: which model backs a
    pipeline is the orchestrator's own choice now, not something this UI
    (or any other caller) gets to set."""

    def __init__(self, parent):
        super().__init__(parent, padding=10)
        self._pipeline_ids = []
        self._pipelines_info = {}
        self.history = []  # [{"role": "user"|"assistant", "content": str}, ...]

        self.conn = ConnectionBar(self, default_port=9000, health_path="/health")
        self.conn.grid(row=0, column=0, columnspan=4, sticky="w", pady=(0, 10))

        ttk.Label(self, text="Pipeline:").grid(row=1, column=0, sticky="w")
        self.pipeline = tk.StringVar()
        self.pipeline_box = ttk.Combobox(self, textvariable=self.pipeline, width=28, state="readonly")
        self.pipeline_box.grid(row=1, column=1, sticky="w")
        self.pipeline_box.bind("<<ComboboxSelected>>", lambda e: self.new_conversation())
        ttk.Button(self, text="Refresh", command=self.refresh_pipelines).grid(row=1, column=2, sticky="w", padx=4)
        ttk.Button(self, text="New conversation", command=self.new_conversation).grid(row=1, column=3, sticky="w")

        self.description = ttk.Label(self, foreground="gray", wraplength=640, justify="left")
        self.description.grid(row=2, column=0, columnspan=4, sticky="w", pady=(4, 0))

        ttk.Label(self, text="Conversation:").grid(row=3, column=0, sticky="nw", pady=(10, 0))
        self.transcript = OutputBox(self, height=16)
        self.transcript.grid(row=4, column=0, columnspan=4, pady=4)

        ttk.Label(self, text="Message:").grid(row=5, column=0, sticky="nw")
        self.message_input = scrolledtext.ScrolledText(self, width=60, height=3, wrap="word")
        self.message_input.grid(row=5, column=1, columnspan=2, sticky="w")
        ttk.Button(self, text="Send", command=self.send).grid(row=5, column=3, sticky="nw")

        ttk.Label(self, text="Result:").grid(row=6, column=0, sticky="nw", pady=(10, 0))
        self.result_box = OutputBox(self, height=8)
        self.result_box.grid(row=7, column=0, columnspan=4, pady=4)

        self.refresh_pipelines()

    def refresh_pipelines(self):
        run_bg(self._refresh_pipelines_bg)

    def _refresh_pipelines_bg(self):
        try:
            r = requests.get(f"{self.conn.base_url}/pipelines", timeout=TIMEOUT_SHORT)
            r.raise_for_status()
            self.after(0, self._apply_pipelines, r.json())
        except requests.RequestException as exc:
            self.after(0, messagebox.showerror, "Orchestrator", f"Could not fetch pipelines: {error_detail(exc)}")

    def _apply_pipelines(self, items):
        self._pipelines_info = {i["id"]: i for i in items}
        self._pipeline_ids = list(self._pipelines_info)
        self.pipeline_box["values"] = [f"{i['id']} - {i['name']}" for i in items]
        if self._pipeline_ids:
            self.pipeline_box.current(0)
            self.new_conversation()

    def _selected_pipeline_id(self):
        idx = self.pipeline_box.current()
        return self._pipeline_ids[idx] if 0 <= idx < len(self._pipeline_ids) else None

    def new_conversation(self):
        pid = self._selected_pipeline_id()
        if pid is None:
            return
        self.description.config(text=self._pipelines_info.get(pid, {}).get("description", ""))
        self.history = []
        self.transcript.write("")
        self.result_box.write("")
        self._send(pid, {"history": []})

    def send(self):
        text = self.message_input.get("1.0", tk.END).strip()
        if not text:
            return
        pid = self._selected_pipeline_id()
        if pid is None:
            messagebox.showwarning("Orchestrator", "Pick a pipeline first.")
            return
        self.message_input.delete("1.0", tk.END)
        self._append_line(f"You: {text}")
        self._send(pid, {"history": self.history, "text": text})

    def _send(self, pipeline_id, payload):
        run_bg(self._send_bg, pipeline_id, payload)

    def _send_bg(self, pipeline_id, payload):
        try:
            r = requests.post(f"{self.conn.base_url}/pipelines/{pipeline_id}/run",
                              json=payload, timeout=TIMEOUT_LOAD)
            r.raise_for_status()
            self.after(0, self._apply_reply, payload, r.json())
        except requests.RequestException as exc:
            self.after(0, self._append_line, f"Error: {error_detail(exc)}")

    def _apply_reply(self, payload, body):
        reply = body.get("reply", "")
        if payload.get("text"):
            self.history.append({"role": "user", "content": payload["text"]})
        if reply:
            self.history.append({"role": "assistant", "content": reply})
            self._append_line(f"Assistant: {reply}")

        status = body.get("status")
        if status:
            self._append_line(f"[{status}]")

        # Whatever the pipeline returned beyond reply/status (e.g. greeting's
        # patient_situation) — shown as-is, since each pipeline's result shape
        # is its own business, not something this generic tab needs to know.
        extra = {k: v for k, v in body.items() if k not in ("reply", "status")}
        if any(v is not None for v in extra.values()):
            self.result_box.write(json.dumps(extra, indent=2, ensure_ascii=False))

    def _append_line(self, line):
        current = self.transcript.get("1.0", tk.END).rstrip("\n")
        self.transcript.write(f"{current}\n{line}" if current else line)


class DemoApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Spin Medical Assistant — Demo")
        self.geometry("900x760")
        # Freely resizable/maximizable -- content that doesn't fit is reached
        # by scrolling (ScrollableFrame below) rather than by capping the
        # window's growth, so there's no need to restrict maxsize here.

        notebook = ttk.Notebook(self)
        notebook.pack(fill="both", expand=True)

        for tab_cls, label in [(STTTab, "STT"), (LLMTab, "Core_LLM"), (OrchestratorTab, "Orchestrator")]:
            scroller = ScrollableFrame(notebook)
            tab_cls(scroller.body).pack(fill="both", expand=True)
            notebook.add(scroller, text=label)


if __name__ == "__main__":
    DemoApp().mainloop()
